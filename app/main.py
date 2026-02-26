import uuid
import io
import re
from fastapi import FastAPI, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.prompts.templates import build_prompt
from app.services.llm_service import generate_article

app = FastAPI(title="SEO Bot — Деликатная Терапия Души")

# In-memory session storage (single user)
sessions: dict[str, list[dict]] = {}

app.mount("/static", StaticFiles(directory="static"), name="static")


class ArticleRequest(BaseModel):
    topic: str
    h1: str = ""
    content_type: str = "article"
    main_keywords: str
    thematic_words: str = ""
    highlight_words: str = ""
    word_count: int = 3000
    structure: str = ""
    aeo_questions: str = ""
    meta_title: str = ""
    meta_description: str = ""
    competitors: str = ""
    additional: str = ""


class ChatMessage(BaseModel):
    session_id: str
    message: str


class BriefTextRequest(BaseModel):
    text: str


class ExportRequest(BaseModel):
    text: str


@app.get("/")
async def index():
    return FileResponse("static/index.html")


@app.post("/api/upload-brief")
async def upload_brief(file: UploadFile = File(...)):
    """Parse an uploaded DOCX/TXT brief and return its text content."""
    content = await file.read()
    filename = file.filename.lower()

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif filename.endswith(".txt"):
        text = content.decode("utf-8")
    else:
        return {"error": "Поддерживаются только .docx и .txt файлы"}

    return {"text": text}


@app.post("/api/generate")
async def generate(req: ArticleRequest):
    """Generate a new article from a brief form."""
    session_id = str(uuid.uuid4())

    user_prompt = build_prompt(
        topic=req.topic,
        h1=req.h1,
        content_type=req.content_type,
        main_keywords=req.main_keywords,
        thematic_words=req.thematic_words,
        highlight_words=req.highlight_words,
        word_count=req.word_count,
        structure=req.structure,
        aeo_questions=req.aeo_questions,
        meta_title=req.meta_title,
        meta_description=req.meta_description,
        competitors=req.competitors,
        additional=req.additional,
    )

    messages = [{"role": "user", "content": user_prompt}]
    sessions[session_id] = messages.copy()

    try:
        result = generate_article(messages)
        sessions[session_id].append({"role": "assistant", "content": result})
        return {"text": result, "session_id": session_id}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/generate-from-brief")
async def generate_from_brief(req: BriefTextRequest):
    """Generate an article from uploaded brief text."""
    session_id = str(uuid.uuid4())

    user_prompt = (
        "Вот ТЗ на статью. Напиши статью строго по этому ТЗ.\n\n"
        "ВАЖНО:\n"
        "- Соблюдай указанную структуру (H2/H3 заголовки)\n"
        "- Используй все ключевые слова в точной форме и указанной частотности\n"
        "- Соблюдай указанный объём в символах без пробелов\n"
        "- Обеспечь логические переходы между всеми разделами\n"
        "- Верни результат в формате ---META--- и ---ARTICLE---\n\n"
        f"{req.text}"
    )

    messages = [{"role": "user", "content": user_prompt}]
    sessions[session_id] = messages.copy()

    try:
        result = generate_article(messages)
        sessions[session_id].append({"role": "assistant", "content": result})
        return {"text": result, "session_id": session_id}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/chat")
async def chat(req: ChatMessage):
    """Continue conversation for article refinement."""
    if req.session_id not in sessions:
        return JSONResponse(status_code=404, content={"error": "Сессия не найдена. Сгенерируйте статью заново."})

    sessions[req.session_id].append({"role": "user", "content": req.message})

    try:
        result = generate_article(sessions[req.session_id])
        sessions[req.session_id].append({"role": "assistant", "content": result})
        return {"text": result}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/export-docx")
async def export_docx(req: ExportRequest):
    """Export article as DOCX file."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    meta_match = re.search(r"---META---([\s\S]*?)---/META---", req.text)
    article_match = re.search(r"---ARTICLE---([\s\S]*?)---/ARTICLE---", req.text)

    # Add meta block
    if meta_match:
        meta_text = meta_match.group(1).strip()
        title_match = re.search(r"Title:\s*(.+)", meta_text)
        desc_match = re.search(r"Description:\s*(.+)", meta_text)

        if title_match:
            p = doc.add_paragraph()
            run = p.add_run(f"Title: {title_match.group(1).strip()}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = None

        if desc_match:
            p = doc.add_paragraph()
            run = p.add_run(f"Description: {desc_match.group(1).strip()}")
            run.bold = True
            run.font.size = Pt(10)

        doc.add_paragraph("")  # spacer

    # Get article text
    article_text = article_match.group(1).strip() if article_match else req.text
    if not article_match:
        article_text = re.sub(r"---META---[\s\S]*?---/META---", "", article_text)
        article_text = article_text.replace("---ARTICLE---", "").replace("---/ARTICLE---", "").strip()

    # Parse markdown and add to doc
    for line in article_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")
        else:
            # Handle bold/italic in paragraph
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=article.docx"},
    )
